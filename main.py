import os
import re
import csv
import tkinter as tk
from tkinter import filedialog, simpledialog, messagebox
from pdf2docx import Converter
from docx import Document
from reportlab.pdfgen import canvas
from PyPDF2 import PdfReader
import comtypes.client
from PIL import Image
from unidecode import unidecode
import subprocess
from odf import text, teletype
from odf.opendocument import load
from bs4 import BeautifulSoup
from striprtf.striprtf import rtf_to_text
import zipfile
import rarfile
import tarfile
import py7zr

CONVERSION_MAP = {
    # Conversões de áudio
    ".mp3": ["wav", "ogg", "flac", "aac", "m4a", "mp4", "avi", "mkv", "webm"],
    ".wav": ["mp3", "ogg", "flac", "aac", "m4a", "mp4", "avi", "mkv", "webm"],
    ".ogg": ["mp3", "wav", "flac", "aac", "m4a", "mp4", "avi", "mkv", "webm"],
    ".flac": ["mp3", "wav", "ogg", "aac", "m4a", "mp4", "avi", "mkv", "webm"],
    ".aac": ["mp3", "wav", "ogg", "flac", "m4a", "mp4", "avi", "mkv", "webm"],
    ".m4a": ["mp3", "wav", "ogg", "flac", "aac", "mp4", "avi", "mkv", "webm"],

    # Conversões de vídeo
    ".mp4": ["avi", "mkv", "webm", "mov", "flv", "wmv", "mp3", "wav", "ogg", "flac", "aac", "m4a"],
    ".avi": ["mp4", "mkv", "webm", "mov", "flv", "wmv", "mp3", "wav", "ogg", "flac", "aac", "m4a"],
    ".mkv": ["mp4", "avi", "webm", "mov", "flv", "wmv", "mp3", "wav", "ogg", "flac", "aac", "m4a"],
    ".webm": ["mp4", "avi", "mkv", "mov", "flv", "wmv", "mp3", "wav", "ogg", "flac", "aac", "m4a"],
    ".mov": ["mp4", "avi", "mkv", "webm", "flv", "wmv", "mp3", "wav", "ogg", "flac", "aac", "m4a"],
    ".flv": ["mp4", "avi", "mkv", "webm", "mov", "wmv", "mp3", "wav", "ogg", "flac", "aac", "m4a"],
    ".wmv": ["mp4", "avi", "mkv", "webm", "mov", "flv", "mp3", "wav", "ogg", "flac", "aac", "m4a"],

    # Conversões de imagem
    ".jpg": ["png", "bmp", "ico", "webp", "jpeg", "pdf", "tiff", "gif"],
    ".jpeg": ["png", "bmp", "ico", "webp", "jpg", "pdf", "tiff", "gif"],
    ".png": ["jpg", "bmp", "ico", "webp", "jpeg", "pdf", "tiff", "gif"],
    ".bmp": ["jpg", "png", "ico", "webp", "jpeg", "pdf", "tiff", "gif"],
    ".ico": ["jpg", "png", "bmp", "webp", "jpeg", "pdf", "tiff", "gif"],
    ".webp": ["jpg", "png", "bmp", "ico", "jpeg", "pdf", "tiff", "gif"],
    ".tiff": ["jpg", "png", "bmp", "ico", "webp", "jpeg", "pdf", "gif"],
    ".gif": ["jpg", "png", "bmp", "ico", "webp", "jpeg", "pdf", "tiff"],

    # Conversões de texto e documentos
    ".txt": ["pdf", "docx", "odt", "html", "rtf", "csv"],
    ".pdf": ["txt", "docx", "jpg", "png", "html", "odt", "rtf"],
    ".docx": ["txt", "pdf", "odt", "html", "rtf", "csv"],
    ".odt": ["txt", "pdf", "docx", "html", "rtf", "csv"],
    ".html": ["txt", "pdf", "docx", "odt", "rtf", "csv"],
    ".rtf": ["txt", "pdf", "docx", "odt", "html", "csv"],
    ".csv": ["txt", "pdf", "docx", "odt", "html", "rtf"],

    # Conversões de outros formatos
    ".zip": ["rar", "tar", "7z"],
    ".rar": ["zip", "tar", "7z"],
    ".tar": ["zip", "rar", "7z"],
    ".7z": ["zip", "rar", "tar"],
}

def sanitize_filename(filename):
    """
    Remove caracteres especiais, acentos e substitui espaços por '-'.
    """
    filename = unidecode(filename)
    filename = re.sub(r'[^\w\-\./]', '', filename)
    filename = re.sub(r'-+/', '', filename)
    filename = filename.strip('')
    return filename

def convert_docx_to_pdf(input_file, output_file):
    try:
        word = comtypes.client.CreateObject("Word.Application")
        doc = word.Documents.Open(input_file)
        doc.SaveAs(output_file, FileFormat=17)
        doc.Close()
        word.Quit()
        return True
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para PDF: {e}")
        return False

def convert_image(input_file, output_file, output_format):
    try:
        img = Image.open(input_file)
        img = img.convert("RGB") if output_format in ["jpeg", "jpg"] else img
        img.save(output_file, format=output_format.upper())
        print(f"✅ Convertido com sucesso para {output_format}: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para {output_format}: {e}")

def convert_video(input_file, output_file):
    try:
        print(f"Convertendo arquivo: {input_file} para {output_file}")

        # Usando subprocess para chamar o FFmpeg diretamente
        command = [
            'ffmpeg', '-i', input_file, '-vcodec', 'libx264', '-acodec', 'aac',
            '-b:v', '2000k', '-f', 'mp4', output_file
        ]

        result = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)

        if result.returncode == 0:
            print(f"✅ Vídeo convertido com sucesso: {output_file}")
        else:
            print(f"❌ Erro ao converter {input_file}: {result.stderr}")

    except Exception as e:
        print(f"❌ Erro ao converter {input_file}: {e}")

def convert_txt_to_html(input_file, output_file):
    try:
        with open(input_file, "r", encoding="utf-8") as f:
            content = f.read()

        html_content = f"<html>\n<body>\n<pre>\n{content}\n</pre>\n</body>\n</html>"

        with open(output_file, "w", encoding="utf-8") as f:
            f.write(html_content)

        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para HTML: {e}")

def convert_txt_to_csv(input_file, output_file, delimiter=','):
    try:
        print(f"Convertendo {input_file} para CSV...")
        with open(input_file, "r", encoding="utf-8") as txt_file:
            lines = txt_file.readlines()

        with open(output_file, "w", encoding="utf-8", newline='') as csv_file:
            writer = csv.writer(csv_file, delimiter=delimiter)
            for line in lines:
                # Remove espaços em branco no início e no fim da linha
                line = line.strip()
                # Divide a linha em colunas com base no delimitador
                columns = line.split()
                writer.writerow(columns)

        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para CSV: {e}")

def convert_docx_to_csv(input_file, output_file, delimiter=','):
    try:
        print(f"Convertendo {input_file} para CSV...")
        doc = Document(input_file)
        with open(output_file, "w", encoding="utf-8", newline='') as csv_file:
            writer = csv.writer(csv_file, delimiter=delimiter)
            for para in doc.paragraphs:
                # Remove espaços em branco no início e no fim da linha
                line = para.text.strip()
                # Divide a linha em colunas com base no delimitador
                columns = line.split()
                writer.writerow(columns)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para CSV: {e}")

def convert_pdf_to_csv(input_file, output_file, delimiter=','):
    try:
        print(f"Convertendo {input_file} para CSV...")
        reader = PdfReader(input_file)
        with open(output_file, "w", encoding="utf-8", newline='') as csv_file:
            writer = csv.writer(csv_file, delimiter=delimiter)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    for line in text.splitlines():
                        # Remove espaços em branco no início e no fim da linha
                        line = line.strip()
                        # Divide a linha em colunas com base no delimitador
                        columns = line.split()
                        writer.writerow(columns)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para CSV: {e}")

def convert_odt_to_csv(input_file, output_file, delimiter=','):
    try:
        print(f"Convertendo {input_file} para CSV...")
        doc = load(input_file)
        with open(output_file, "w", encoding="utf-8", newline='') as csv_file:
            writer = csv.writer(csv_file, delimiter=delimiter)
            for para in doc.getElementsByType(text.P):
                line = teletype.extractText(para).strip()
                columns = line.split()
                writer.writerow(columns)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para CSV: {e}")

def convert_html_to_csv(input_file, output_file, delimiter=','):
    try:
        print(f"Convertendo {input_file} para CSV...")
        with open(input_file, "r", encoding="utf-8") as html_file:
            soup = BeautifulSoup(html_file, "html.parser")
            text = soup.get_text()
        with open(output_file, "w", encoding="utf-8", newline='') as csv_file:
            writer = csv.writer(csv_file, delimiter=delimiter)
            for line in text.splitlines():
                line = line.strip()
                columns = line.split()
                writer.writerow(columns)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para CSV: {e}")

def convert_rtf_to_csv(input_file, output_file, delimiter=','):
    try:
        print(f"Convertendo {input_file} para CSV...")
        with open(input_file, "r", encoding="utf-8") as rtf_file:
            rtf_content = rtf_file.read()
            text = rtf_to_text(rtf_content)
        with open(output_file, "w", encoding="utf-8", newline='') as csv_file:
            writer = csv.writer(csv_file, delimiter=delimiter)
            for line in text.splitlines():
                line = line.strip()
                columns = line.split()
                writer.writerow(columns)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para CSV: {e}")



def convert_csv_to_txt(input_file, output_file):
    try:
        with open(input_file, "r", encoding="utf-8") as csv_file:
            csv_reader = csv.reader(csv_file)
            with open(output_file, "w", encoding="utf-8") as txt_file:
                for row in csv_reader:
                    txt_file.write("\t".join(row) + "\n")
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para TXT: {e}")

def convert_csv_to_pdf(input_file, output_file):
    try:
        c = canvas.Canvas(output_file)
        with open(input_file, "r", encoding="utf-8") as csv_file:
            csv_reader = csv.reader(csv_file)
            y = 800
            for row in csv_reader:
                c.drawString(100, y, "\t".join(row))
                y -= 15
        c.save()
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para PDF: {e}")

def convert_csv_to_docx(input_file, output_file):
    try:
        doc = Document()
        with open(input_file, "r", encoding="utf-8") as csv_file:
            csv_reader = csv.reader(csv_file)
            for row in csv_reader:
                doc.add_paragraph("\t".join(row))
        doc.save(output_file)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para DOCX: {e}")

def convert_csv_to_html(input_file, output_file):
    try:
        with open(input_file, "r", encoding="utf-8") as csv_file:
            csv_reader = csv.reader(csv_file)
            html_content = "<html>\n<body>\n<table>\n"
            for row in csv_reader:
                html_content += "<tr>\n"
                for cell in row:
                    html_content += f"<td>{cell}</td>\n"
                html_content += "</tr>\n"
            html_content += "</table>\n</body>\n</html>"

        with open(output_file, "w", encoding="utf-8") as f:
            f.write(html_content)

        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para HTML: {e}")

def convert_csv_to_rtf(input_file, output_file):
    try:
        with open(input_file, "r", encoding="utf-8") as csv_file:
            csv_reader = csv.reader(csv_file)
            with open(output_file, "w", encoding="utf-8") as rtf_file:
                rtf_file.write("{\\rtf1\\ansi\\deff0\n")
                for row in csv_reader:
                    rtf_file.write("\\trowd\n")
                    for cell in row:
                        rtf_file.write(f"\\cellx{1000}\\intbl {cell}\\cell\n")
                    rtf_file.write("\\row\n")
                rtf_file.write("}")
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para RTF: {e}")

def convert_odt_to_txt(input_file, output_file):
    try:
        doc = load(input_file)
        all_text = []
        for paragraph in doc.getElementsByType(text.P):
            all_text.append(teletype.extractText(paragraph))
        with open(output_file, "w", encoding="utf-8") as f:
            f.write("\n".join(all_text))
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para TXT: {e}")

def convert_odt_to_pdf(input_file, output_file):
    try:
        # Usando o LibreOffice via subprocess para converter ODT para PDF
        command = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(output_file), input_file]
        subprocess.run(command, check=True)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para PDF: {e}")

def convert_odt_to_docx(input_file, output_file):
    try:
        # Usando o LibreOffice via subprocess para converter ODT para DOCX
        command = ['libreoffice', '--headless', '--convert-to', 'docx', '--outdir', os.path.dirname(output_file), input_file]
        subprocess.run(command, check=True)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para DOCX: {e}")

def convert_odt_to_html(input_file, output_file):
    try:
        # Usando o LibreOffice via subprocess para converter ODT para HTML
        command = ['libreoffice', '--headless', '--convert-to', 'html', '--outdir', os.path.dirname(output_file), input_file]
        subprocess.run(command, check=True)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para HTML: {e}")

def convert_odt_to_rtf(input_file, output_file):
    try:
        # Usando o LibreOffice via subprocess para converter ODT para RTF
        command = ['libreoffice', '--headless', '--convert-to', 'rtf', '--outdir', os.path.dirname(output_file), input_file]
        subprocess.run(command, check=True)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para RTF: {e}")

def convert_html_to_txt(input_file, output_file):
    try:
        with open(input_file, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f, "html.parser")
            text = soup.get_text()
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para TXT: {e}")

def convert_html_to_pdf(input_file, output_file):
    try:
        # Usando o wkhtmltopdf para converter HTML para PDF
        command = ['wkhtmltopdf', input_file, output_file]
        subprocess.run(command, check=True)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para PDF: {e}")

def convert_html_to_docx(input_file, output_file):
    try:
        # Usando o pandoc para converter HTML para DOCX
        command = ['pandoc', input_file, '-o', output_file]
        subprocess.run(command, check=True)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para DOCX: {e}")

def convert_html_to_odt(input_file, output_file):
    try:
        # Usando o pandoc para converter HTML para ODT
        command = ['pandoc', input_file, '-o', output_file]
        subprocess.run(command, check=True)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para ODT: {e}")

def convert_html_to_rtf(input_file, output_file):
    try:
        # Usando o pandoc para converter HTML para RTF
        command = ['pandoc', input_file, '-o', output_file]
        subprocess.run(command, check=True)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para RTF: {e}")

def convert_rtf_to_txt(input_file, output_file):
    try:
        with open(input_file, "r", encoding="utf-8") as f:
            rtf_content = f.read()
        text = rtf_to_text(rtf_content)
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para TXT: {e}")

def convert_rtf_to_pdf(input_file, output_file):
    try:
        # Usando o LibreOffice via subprocess para converter RTF para PDF
        command = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(output_file), input_file]
        subprocess.run(command, check=True)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para PDF: {e}")

def convert_rtf_to_docx(input_file, output_file):
    try:
        # Usando o LibreOffice via subprocess para converter RTF para DOCX
        command = ['libreoffice', '--headless', '--convert-to', 'docx', '--outdir', os.path.dirname(output_file), input_file]
        subprocess.run(command, check=True)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para DOCX: {e}")

def convert_rtf_to_odt(input_file, output_file):
    try:
        # Usando o LibreOffice via subprocess para converter RTF para ODT
        command = ['libreoffice', '--headless', '--convert-to', 'odt', '--outdir', os.path.dirname(output_file), input_file]
        subprocess.run(command, check=True)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para ODT: {e}")

def convert_rtf_to_html(input_file, output_file):
    try:
        # Usando o pandoc para converter RTF para HTML
        command = ['pandoc', input_file, '-o', output_file]
        subprocess.run(command, check=True)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para HTML: {e}")

def convert_zip_to_rar(input_file, output_file):
    try:
        with zipfile.ZipFile(input_file, 'r') as zip_ref:
            with rarfile.RarFile(output_file, 'w') as rar_ref:
                for file in zip_ref.namelist():
                    rar_ref.write(file)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para RAR: {e}")

def convert_zip_to_tar(input_file, output_file):
    try:
        with zipfile.ZipFile(input_file, 'r') as zip_ref:
            with tarfile.open(output_file, 'w') as tar_ref:
                for file in zip_ref.namelist():
                    tar_ref.add(file)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para TAR: {e}")

def convert_zip_to_7z(input_file, output_file):
    try:
        with zipfile.ZipFile(input_file, 'r') as zip_ref:
            with py7zr.SevenZipFile(output_file, 'w') as seven_zip_ref:
                for file in zip_ref.namelist():
                    seven_zip_ref.write(file)
        print(f"✅ Convertido com sucesso: {output_file}")
    except Exception as e:
        print(f"❌ Erro ao converter {input_file} para 7Z: {e}")

def convert_file(input_files, output_format):
    output_dir = os.path.join(os.path.expanduser("~"), "converted_files")
    os.makedirs(output_dir, exist_ok=True)

    for file in input_files:
        base_name, ext = os.path.splitext(os.path.basename(file))
        base_name = sanitize_filename(base_name)
        output_file = os.path.join(output_dir, f"{base_name}.{output_format}")

        try:
            if ext in [".mp3", ".wav", ".mp4", ".avi"]:
                convert_video(file, output_file)
            elif ext == ".pdf" and output_format == "docx":
                cv = Converter(file)
                cv.convert(output_file, start=0, end=None)
                cv.close()
            elif ext == ".pdf" and output_format == "txt":
                reader = PdfReader(file)
                text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
                with open(output_file, "w", encoding="utf-8") as f:
                    f.write(text)
            elif ext == ".docx" and output_format == "txt":
                doc = Document(file)
                text = "\n".join([para.text for para in doc.paragraphs])
                with open(output_file, "w", encoding="utf-8") as f:
                    f.write(text)
            elif ext == ".docx" and output_format == "pdf":
                if convert_docx_to_pdf(file, output_file):
                    print(f"✅ Convertido com sucesso: {output_file}")
                else:
                    print(f"❌ Erro: O arquivo {output_file} não foi gerado!")
            elif ext == ".txt" and output_format == "docx":
                doc = Document()
                with open(file, "r", encoding="utf-8") as f:
                    for line in f:
                        doc.add_paragraph(line)
                doc.save(output_file)
            elif ext == ".txt" and output_format == "pdf":
                c = canvas.Canvas(output_file)
                with open(file, "r", encoding="utf-8") as f:
                    text = f.readlines()
                y = 800
                for line in text:
                    c.drawString(100, y, line.strip())
                    y -= 15
                c.save()
            elif ext == ".txt" and output_format == "html":
                convert_txt_to_html(file, output_file)
            elif ext == ".txt" and output_format == "csv":
                convert_txt_to_csv(file, output_file)
            elif ext == ".csv" and output_format == "txt":
                convert_csv_to_txt(file, output_file)
            elif ext == ".csv" and output_format == "pdf":
                convert_csv_to_pdf(file, output_file)
            elif ext == ".csv" and output_format == "docx":
                convert_csv_to_docx(file, output_file)
            elif ext == ".csv" and output_format == "html":
                convert_csv_to_html(file, output_file)
            elif ext == ".csv" and output_format == "rtf":
                convert_csv_to_rtf(file, output_file)
            elif ext == ".txt" and output_format == "csv":
                convert_txt_to_csv(file, output_file)
            elif ext == ".docx" and output_format == "csv":
                convert_docx_to_csv(file, output_file)
            elif ext == ".pdf" and output_format == "csv":
                convert_pdf_to_csv(file, output_file)
            elif ext == ".odt" and output_format == "csv":
                convert_odt_to_csv(file, output_file)
            elif ext == ".html" and output_format == "csv":
                convert_html_to_csv(file, output_file)
            elif ext == ".rtf" and output_format == "csv":
                convert_rtf_to_csv(file, output_file)
            elif ext == ".csv" and output_format == "txt":
                convert_csv_to_txt(file, output_file)
            elif ext == ".csv" and output_format == "pdf":
                convert_csv_to_pdf(file, output_file)
            elif ext == ".csv" and output_format == "docx":
                convert_csv_to_docx(file, output_file)
            elif ext == ".csv" and output_format == "html":
                convert_csv_to_html(file, output_file)
            elif ext == ".csv" and output_format == "rtf":
                convert_csv_to_rtf(file, output_file)
            elif ext == ".odt" and output_format == "txt":
                convert_odt_to_txt(file, output_file)
            elif ext == ".odt" and output_format == "pdf":
                convert_odt_to_pdf(file, output_file)
            elif ext == ".odt" and output_format == "docx":
                convert_odt_to_docx(file, output_file)
            elif ext == ".odt" and output_format == "html":
                convert_odt_to_html(file, output_file)
            elif ext == ".odt" and output_format == "rtf":
                convert_odt_to_rtf(file, output_file)
            elif ext == ".html" and output_format == "txt":
                convert_html_to_txt(file, output_file)
            elif ext == ".html" and output_format == "pdf":
                convert_html_to_pdf(file, output_file)
            elif ext == ".html" and output_format == "docx":
                convert_html_to_docx(file, output_file)
            elif ext == ".html" and output_format == "odt":
                convert_html_to_odt(file, output_file)
            elif ext == ".html" and output_format == "rtf":
                convert_html_to_rtf(file, output_file)
            elif ext == ".rtf" and output_format == "txt":
                convert_rtf_to_txt(file, output_file)
            elif ext == ".rtf" and output_format == "pdf":
                convert_rtf_to_pdf(file, output_file)
            elif ext == ".rtf" and output_format == "docx":
                convert_rtf_to_docx(file, output_file)
            elif ext == ".rtf" and output_format == "odt":
                convert_rtf_to_odt(file, output_file)
            elif ext == ".rtf" and output_format == "html":
                convert_rtf_to_html(file, output_file)
            elif ext == ".zip" and output_format == "rar":
                convert_zip_to_rar(file, output_file)
            elif ext == ".zip" and output_format == "tar":
                convert_zip_to_tar(file, output_file)
            elif ext == ".zip" and output_format == "7z":
                convert_zip_to_7z(file, output_file)
            elif ext in [".jpg", ".png", ".bmp", ".webp"]:
                convert_image(file, output_file, output_format)

            if os.path.exists(output_file):
                print(f"✅ Convertido com sucesso: {output_file}")
                # Exibe uma notificação para cada arquivo convertido
                messagebox.showinfo("Conversão concluída", f"Arquivo convertido com sucesso:\n{output_file}")
            else:
                print(f"❌ Erro: O arquivo {output_file} não foi gerado!")
        except Exception as e:
            print(f"❌ Erro ao converter {file}: {e}")

def select_files():
    files = filedialog.askopenfilenames(title="Selecione arquivos")
    if not files:
        return

    ext = os.path.splitext(files[0])[1].lower()
    options = CONVERSION_MAP.get(ext, [])

    if not options:
        print(f"⚠️ Não há opções de conversão disponíveis para {ext}")
        return

    chosen_format = simpledialog.askstring("Escolha o formato", f"Opções para {ext}: {', '.join(options)}\nDigite o formato desejado:")

    if chosen_format in options:
        convert_file(files, chosen_format)
    else:
        print(f"⚠️ Formato inválido: {chosen_format}")

root = tk.Tk()
root.withdraw()

select_files()